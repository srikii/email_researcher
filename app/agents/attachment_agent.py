from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.graph.state import EmailResearchState
from app.models import ExtractedDocument


def attachment_agent(state: EmailResearchState) -> EmailResearchState:
    # This agent turns downloaded files into plain text documents.
    documents: list[ExtractedDocument] = []
    errors: list[str] = []

    for email in state.get("emails", []):
        for attachment in email.attachments:
            # Gmail metadata can say an attachment exists before we have
            # downloaded it. If local_path is missing, there is no file to read.
            if not attachment.local_path:
                continue

            path = Path(attachment.local_path)
            try:
                text = extract_attachment_text(path)
                if text.strip():
                    # Keep enough metadata so search results can point back to
                    # the email that contained this attachment.
                    documents.append(
                        ExtractedDocument(
                            source_type="attachment",
                            source=str(path),
                            title=attachment.filename,
                            text=text,
                            email_id=email.id,
                            metadata={
                                "subject": email.subject,
                                "sender": email.sender,
                                "date": email.date,
                                "mime_type": attachment.mime_type,
                            },
                        )
                    )
            except Exception as exc:
                errors.append(f"attachment_agent failed for {attachment.filename}: {exc}")

    return {"documents": documents, "errors": errors}


def extract_attachment_text(path: Path) -> str:
    # Choose the parser based on the file extension.
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        page_texts: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            page_texts.append(page_text)
        return "\n\n".join(page_texts)

    if suffix in {".docx", ".doc"}:
        doc = Document(str(path))
        paragraph_texts: list[str] = []
        for paragraph in doc.paragraphs:
            paragraph_texts.append(paragraph.text)
        return "\n".join(paragraph_texts)

    raise ValueError(f"Unsupported attachment type: {path.suffix}")
